"""Export per-fund MtM transaction histories from the 2025 Excel sheet to CSV
and generate Form 8621 Acquisition Date Statement docx files.

Data source sheet: 2025年MtM交易记录（仅年底仍持有基金+多次赎回）
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from qieman_parser.utils import get_taxes_2018_root_path

EXCEL_PATH = r"C:\Users\kunw\OneDrive - KUNW-MSFT\2025Tax\ForeignAccountDetails.xlsx"
SHEET_NAME = "2025年MtM交易记录（仅年底仍持有基金+多次赎回）"

FONT_NAME = "Courier New"
FONT_SIZE = Pt(11)

ACCOUNT_TO_ENGLISH_MAP = {
    "雪球": "XQ",
    "涨乐": "ZL",
    "支付宝": "ZFB",
    "国投瑞银": "GTRY",
    "中欧": "ZO",
    "理财通": "LCT",
    "且慢": "QM",
    "招商银行": "ZSYH",
}


def read_usd_to_cny_rate() -> float:
    """Read the USD/CNY exchange rate from cell AL75 in 人民币基金详情."""
    rate_df = pd.read_excel(EXCEL_PATH, sheet_name="人民币基金详情", header=None, nrows=75)
    return float(rate_df.iloc[74, 37])


def _get_reference_name(row: pd.Series) -> str:
    if pd.isna(row.get("基金编号", None)) or str(row["基金编号"]) == "nan":
        pfic_name = row["自定基金编号"]
    else:
        pfic_name = str(row["基金编号"])
    return ACCOUNT_TO_ENGLISH_MAP[row["账号"]] + "." + pfic_name


def build_fund_lookup(
    df_cny: pd.DataFrame, df_usd: pd.DataFrame
) -> dict[str, dict]:
    """Map transaction-sheet fund codes to detail-sheet metadata.

    Returns: {fund_code: {"english_name": ..., "reference_id": ...,
                          "shares_at_start_2025": ..., "is_usd": bool}}
    """
    lookup: dict[str, dict] = {}

    for detail_df, is_usd in [(df_cny, False), (df_usd, True)]:
        shares_col = "2024年底持有份额"
        for _, row in detail_df.iterrows():
            if pd.isna(row["账号"]) or row["账号"] not in ACCOUNT_TO_ENGLISH_MAP:
                continue

            code_from_bianhao = str(row["基金编号"]) if pd.notna(row["基金编号"]) else None
            code_from_ziding = str(row["自定基金编号"]) if pd.notna(row["自定基金编号"]) else None

            ref_id = _get_reference_name(row)
            info = {
                "english_name": row["基金英文"],
                "reference_id": ref_id,
                "shares_at_start_2025": row[shares_col] if pd.notna(row.get(shares_col)) else 0,
                "is_usd": is_usd,
            }

            if code_from_bianhao and code_from_bianhao != "nan":
                lookup[code_from_bianhao] = info
            if code_from_ziding and code_from_ziding != "nan":
                suffix = code_from_ziding.rsplit(".", 1)[-1] if "." in code_from_ziding else code_from_ziding
                lookup[suffix] = info

    return lookup


def _set_run(paragraph, text, bold=False):
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold
    return run


def _set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    _set_run(p, text, bold=bold)


def generate_docx(
    fund_code: str,
    fund_info: dict,
    txn_df: pd.DataFrame,
    output_path: str,
):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE

    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(title_para, "Form 8621 Acquisition Date Statement")

    name_para = doc.add_paragraph()
    _set_run(name_para, fund_info["english_name"])

    ref_para = doc.add_paragraph()
    _set_run(ref_para, fund_info["reference_id"])

    doc.add_paragraph()

    shares = fund_info["shares_at_start_2025"]
    if isinstance(shares, float) and shares == int(shares):
        shares_str = str(int(shares))
    else:
        shares_str = f"{shares:.2f}" if isinstance(shares, (int, float)) else str(shares)
    shares_para = doc.add_paragraph()
    _set_run(shares_para, f"At the beginning of 2025, {shares_str} shares were held.")

    doc.add_paragraph()

    is_usd = fund_info["is_usd"]
    if is_usd:
        headers = ["Date", "Transaction type", "Number of shares", "Amount (USD)"]
    else:
        headers = ["Date", "Transaction type", "Number of shares", "Amount (USD)"]

    table = doc.add_table(rows=1, cols=4, style="Table Grid")

    for ci, header_text in enumerate(headers):
        _set_cell_text(table.rows[0].cells[ci], header_text, bold=True)

    total_shares = 0.0
    total_amount = 0.0

    for _, txn in txn_df.iterrows():
        row_cells = table.add_row().cells

        date_val = txn["交易日期"]
        try:
            date_str = f"{date_val.month}/{date_val.day}/{date_val.year}"
        except AttributeError:
            date_str = str(date_val)

        shares_val = txn["份额"]
        if is_usd:
            amount_val = txn.get("金额（US$）", 0)
        else:
            amount_val = txn.get("金额（US$）", None)
            if pd.isna(amount_val):
                amount_val = txn.get("金额（人民币）", 0)

        _set_cell_text(row_cells[0], date_str)
        _set_cell_text(row_cells[1], str(txn["操作"]))
        _set_cell_text(row_cells[2], f"{shares_val:.2f}" if pd.notna(shares_val) else "")
        _set_cell_text(row_cells[3], f"{amount_val:.2f}" if pd.notna(amount_val) else "")

        if pd.notna(shares_val):
            total_shares += float(shares_val)
        if pd.notna(amount_val):
            total_amount += float(amount_val)

    total_cells = table.add_row().cells
    _set_cell_text(total_cells[0], "Total", bold=True)
    _set_cell_text(total_cells[1], "")
    _set_cell_text(total_cells[2], f"{total_shares:.2f}", bold=True)
    _set_cell_text(total_cells[3], f"{total_amount:.2f}", bold=True)

    doc.save(output_path)


if __name__ == "__main__":
    repo_root = get_taxes_2018_root_path()
    output_dir = os.path.join(repo_root, "filled", "2025", "transactions")
    os.makedirs(output_dir, exist_ok=True)

    usd_to_cny = read_usd_to_cny_rate()
    print(f"Using USD/CNY exchange rate: {usd_to_cny}")

    df_cny = pd.read_excel(EXCEL_PATH, sheet_name="人民币基金详情")
    df_usd = pd.read_excel(EXCEL_PATH, sheet_name="美元基金详情")
    fund_lookup = build_fund_lookup(df_cny, df_usd)

    df = pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME)

    for fund_name in df["基金名称"].unique():
        fund_df = df[df["基金名称"] == fund_name].copy()
        fund_df = fund_df.sort_values("交易日期")

        is_cny_fund = fund_df["人民币基金每份额价格"].notna().any()
        is_usd_fund = fund_df["美元基金每份额价格"].notna().any()

        if is_cny_fund:
            out_df = pd.DataFrame({
                "Date": fund_df["交易日期"],
                "Transaction type": fund_df["操作"],
                "Cost per share (CNY)": fund_df["人民币基金每份额价格"],
                "Amount (CNY)": fund_df["金额（人民币）"],
                "Share": fund_df["份额"],
                "Amount (USD)": fund_df["金额（US$）"] if fund_df["金额（US$）"].notna().any()
                    else fund_df["金额（人民币）"] / usd_to_cny,
            })
        elif is_usd_fund:
            out_df = pd.DataFrame({
                "Date": fund_df["交易日期"],
                "Transaction type": fund_df["操作"],
                "Cost per share (USD)": fund_df["美元基金每份额价格"],
                "Amount (USD)": fund_df["金额（US$）"],
                "Share": fund_df["份额"],
            })
        else:
            print(f"  Skipping {fund_name}: no price-per-share data")
            continue

        csv_path = os.path.join(output_dir, f"{fund_name}.csv")
        out_df.to_csv(csv_path, index=False, float_format="%.2f")
        print(f"  {fund_name}: {len(out_df)} transactions -> {csv_path}")

        if fund_name not in fund_lookup:
            print(f"    WARNING: {fund_name} not found in detail sheets, skipping docx")
            continue

        fund_info = fund_lookup[fund_name]
        docx_path = os.path.join(output_dir, f"{fund_name}.AcquisitionDateStatement.docx")
        generate_docx(fund_name, fund_info, fund_df, docx_path)
        print(f"    -> {docx_path}")

    print(f"\nDone. Output in {output_dir}")
